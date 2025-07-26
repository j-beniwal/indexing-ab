import os
import json
import yaml
import fitz  # PyMuPDF
import pytesseract
import docx
import tempfile
from PIL import Image
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from llama_index.core import Document, VectorStoreIndex
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.gemini import Gemini
from llama_index.core.settings import Settings
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from pydantic import BaseModel

# Environment variables
EMBEDDING_MODEL_NAME = os.getenv("EMBEDDING_MODEL_NAME", "BAAI/bge-large-en")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
CHROMA_HOST = os.getenv("CHROMA_HOST", "localhost")
CHROMA_PORT = int(os.getenv("CHROMA_PORT", 8000))
CHROMA_COLLECTION_NAME = os.getenv("CHROMA_COLLECTION_NAME", "rag-compliance")
CHROMA_PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./chroma_data")

# Initialize models
embedding_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL_NAME)
llm = Gemini(api_key=GEMINI_API_KEY, model_name="gemini-1.5-flash")

# Global settings
Settings.llm = llm
Settings.embed_model = embedding_model

# PII tools
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# FastAPI app
app = FastAPI()

def redact_text(text: str) -> str:
    analyzer_results = analyzer.analyze(text=text, language='en')
    anonymized_result = anonymizer.anonymize(text=text, analyzer_results=analyzer_results)
    return anonymized_result.text

def extract_documents(file_path: str, filename: str):
    documents = []
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            with fitz.open(file_path) as doc:
                for page_number, page in enumerate(doc, start=1):
                    text = page.get_text()
                    if not text.strip():
                        continue
                    redacted = redact_text(text)
                    documents.append(Document(text=redacted, metadata={"source_file": filename, "page_number": page_number}))

        elif ext == ".docx":
            doc_file = docx.Document(file_path)
            for i, para in enumerate(doc_file.paragraphs):
                text = para.text
                if not text.strip():
                    continue
                redacted = redact_text(text)
                documents.append(Document(text=redacted, metadata={"source_file": filename, "block_number": i + 1}))

        elif ext in [".png", ".jpg", ".jpeg"]:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            if text.strip():
                redacted = redact_text(text)
                documents.append(Document(text=redacted, metadata={"source_file": filename}))

        elif ext in [".yaml", ".yml"]:
            with open(file_path, 'r') as f:
                data = yaml.safe_load(f)
            redacted = redact_text(json.dumps(data))
            documents.append(Document(text=redacted, metadata={"source_file": filename}))

        elif ext == ".json":
            with open(file_path, 'r') as f:
                data = json.load(f)
            redacted = redact_text(json.dumps(data))
            documents.append(Document(text=redacted, metadata={"source_file": filename}))

        elif ext == ".txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                for i, line in enumerate(f):
                    if not line.strip():
                        continue
                    redacted = redact_text(line)
                    documents.append(Document(text=redacted, metadata={"source_file": filename, "line_number": i + 1}))
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {str(e)}")

    if not documents:
        raise HTTPException(status_code=400, detail="No valid content extracted.")

    return documents

@app.post("/ingest_and_query")
async def ingest_and_query(file: UploadFile = File(...), query: str = Form(...)):
    try:
        with tempfile.NamedTemporaryFile(delete=False) as tmp:
            tmp.write(await file.read())
            temp_path = tmp.name

        documents = extract_documents(temp_path, file.filename)

        vector_store = ChromaVectorStore(
            collection_name=CHROMA_COLLECTION_NAME,
            persist_dir=CHROMA_PERSIST_DIR,
            chroma_db_impl="duckdb+parquet",
            host=CHROMA_HOST,
            port=CHROMA_PORT
        )

        index = VectorStoreIndex.from_documents(documents, vector_store=vector_store)
        index.storage_context.persist()

        query_engine = index.as_query_engine(similarity_top_k=5)
        response = query_engine.query(query)

        sources = []
        for node in response.source_nodes:
            metadata = node.metadata or {}
            sources.append({
                "file": metadata.get("source_file", "unknown"),
                "page": metadata.get("page_number", "N/A"),
                "block": metadata.get("block_number", metadata.get("line_number", "N/A")),
                "score": node.score
            })

        os.remove(temp_path)

        return {
            "answer": str(response),
            "sources": sources
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error during processing: {str(e)}")

@app.get("/")
def root():
    return {"message": "RAG ingestion + search API running"}
